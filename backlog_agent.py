"""
backlog_agent.py
─────────────────────────────────────────────────────────────
Agente de criação de backlog Scrum a partir de documentos de requisitos.

Uso:
    python backlog_agent.py

Dependências:
    pip install anthropic python-docx openpyxl pymupdf python-dotenv

Variáveis de ambiente (.env):
    ANTHROPIC_API_KEY=sk-ant-...

Modelos disponíveis (variável MODEL — linha 36):
    claude-haiku-4-5-20251001   → testes (mais barato)
    claude-sonnet-4-6           → produção (melhor qualidade)
"""

import os, json, re, textwrap, datetime
from pathlib import Path
from dotenv import load_dotenv
import anthropic

# ── Imports de geração de arquivos ───────────────────────────
from docx import Document as DocxDocument
from docx.shared import Pt, RGBColor, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
import openpyxl
from openpyxl.styles import Font, PatternFill, Alignment, Border, Side
from openpyxl.utils import get_column_letter
from openpyxl.worksheet.datavalidation import DataValidation

load_dotenv()
CLIENT = anthropic.Anthropic(api_key=os.getenv("ANTHROPIC_API_KEY"))
MODEL = "claude-haiku-4-5-20251001"

# ════════════════════════════════════════════════════════════
# UTILITÁRIOS DE TERMINAL
# ════════════════════════════════════════════════════════════

AZUL    = "\033[94m"
VERDE   = "\033[92m"
AMARELO = "\033[93m"
CINZA   = "\033[90m"
RESET   = "\033[0m"
BOLD    = "\033[1m"

def titulo(t):   print(f"\n{BOLD}{AZUL}{'═'*60}\n  {t}\n{'═'*60}{RESET}")
def secao(t):    print(f"\n{BOLD}{VERDE}── {t}{RESET}")
def info(t):     print(f"{CINZA}   {t}{RESET}")
def pergunta(t): return input(f"\n{AMARELO}▶  {t}{RESET} ").strip()
def ok(t):       print(f"{VERDE}   ✓ {t}{RESET}")
def aviso(t):    print(f"{AMARELO}   ⚠ {t}{RESET}")

# ════════════════════════════════════════════════════════════
# LEITURA DE DOCUMENTOS
# ════════════════════════════════════════════════════════════

def ler_docx(caminho: str) -> str:
    from docx import Document
    doc = Document(caminho)
    return "\n".join(p.text for p in doc.paragraphs if p.text.strip())

def ler_pdf(caminho: str) -> str:
    import fitz
    doc = fitz.open(caminho)
    return "\n".join(page.get_text() for page in doc)

def ler_documento(caminho: str) -> str:
    ext = Path(caminho).suffix.lower()
    if ext == ".docx": return ler_docx(caminho)
    if ext == ".pdf":  return ler_pdf(caminho)
    if ext == ".txt":
        return Path(caminho).read_text(encoding="utf-8")
    raise ValueError(f"Formato não suportado: {ext}. Use .docx, .pdf ou .txt")

def coletar_documentos() -> str:
    """Coleta um ou mais documentos do usuário e retorna o texto concatenado."""
    secao("Documentos de Entrada")
    info("Informe os caminhos dos documentos (ENTER em branco para finalizar).")
    info("Formatos aceitos: .docx · .pdf · .txt")
    textos = []
    idx = 1
    while True:
        caminho = pergunta(f"Documento {idx} (ou ENTER para continuar):").strip('"').strip("'")
        if not caminho:
            if not textos:
                aviso("Nenhum documento informado. Tente novamente.")
                continue
            break
        if not Path(caminho).exists():
            aviso(f"Arquivo não encontrado: {caminho}")
            continue
        try:
            texto = ler_documento(caminho)
            textos.append(f"=== DOCUMENTO {idx}: {Path(caminho).name} ===\n{texto}")
            ok(f"Lido: {Path(caminho).name} ({len(texto):,} caracteres)")
            idx += 1
        except Exception as e:
            aviso(f"Erro ao ler documento: {e}")
    return "\n\n".join(textos)

def coletar_mudancas() -> str:
    """Coleta mudanças/melhorias via terminal ou arquivo .txt."""
    secao("Entrada de Mudanças / Melhorias")
    info("Como deseja informar as mudanças?")
    info("  [1] Digitar diretamente no terminal")
    info("  [2] Informar um arquivo .txt com as mudanças listadas")

    while True:
        opcao = pergunta("Escolha (1 ou 2):").strip()
        if opcao in ("1", "2"):
            break
        aviso("Opção inválida. Digite 1 ou 2.")

    if opcao == "2":
        while True:
            caminho = pergunta("Caminho do arquivo .txt:").strip('"').strip("'")
            if not caminho:
                aviso("Informe o caminho do arquivo.")
                continue
            p = Path(caminho)
            if not p.exists():
                aviso(f"Arquivo não encontrado: {caminho}")
                continue
            if p.suffix.lower() != ".txt":
                aviso("Apenas arquivos .txt são aceitos neste modo.")
                continue
            texto = p.read_text(encoding="utf-8")
            ok(f"Arquivo lido: {p.name} ({len(texto):,} caracteres)")
            return texto

    # Modo terminal: digitar mudanças linha a linha
    info("Digite cada mudança/melhoria e pressione ENTER.")
    info("Seja específico: descreva O QUE muda e POR QUÊ (se souber).")
    info("Para finalizar, pressione ENTER numa linha em branco.\n")
    mudancas = []
    idx = 1
    while True:
        item = pergunta(f"Mudança {idx}:").strip()
        if not item:
            if not mudancas:
                aviso("Nenhuma mudança informada. Tente novamente.")
                continue
            break
        mudancas.append(f"{idx}. {item}")
        ok(f"Registrado: {item[:80]}{'...' if len(item) > 80 else ''}")
        idx += 1
    return "\n".join(mudancas)

# ════════════════════════════════════════════════════════════
# EXTRAÇÃO DE REQUISITOS VIA CLAUDE
# ════════════════════════════════════════════════════════════

PROMPT_EXTRACAO_BASE = """
Você é um analista de requisitos sênior especializado em Scrum e desenvolvimento low code.

Analise o(s) documento(s) abaixo e extraia TODAS as informações relevantes para construção de um backlog Scrum.

Retorne APENAS um JSON válido, sem texto adicional, sem markdown, sem backticks, com esta estrutura exata:

{{
  "projeto": "nome do projeto",
  "sistema": "nome do sistema/plataforma",
  "descricao": "descrição resumida do projeto em 2-3 linhas",
  "atores": ["lista de perfis/atores identificados"],
  "epicos": [
    {{
      "id": "EP01",
      "nome": "nome do épico",
      "descricao": "descrição do épico"
    }}
  ],
  "requisitos_funcionais": [
    {{
      "id": "RF01",
      "descricao": "descrição completa do requisito",
      "epico_id": "EP01",
      "atores": ["Secretaria"]
    }}
  ],
  "regras_negocio": [
    {{
      "id": "RN01",
      "descricao": "descrição da regra"
    }}
  ],
  "requisitos_nao_funcionais": [
    {{
      "id": "RNF01",
      "descricao": "descrição do requisito não funcional"
    }}
  ],
  "fora_de_escopo": ["item 1", "item 2"],
  "dependencias_tecnicas": ["dependência 1 → dependência 2"],
  "observacoes": "observações relevantes sobre o projeto"
}}

DOCUMENTOS:
{documentos}
"""

PROMPT_GERAR_US_BASE = """
Você é um analista de requisitos sênior especializado em Scrum.

Com base nos requisitos extraídos abaixo, gere as Histórias de Usuário (User Stories) para o backlog.

Regras obrigatórias:
- Cada US deve cobrir exatamente 1 entrega funcional testável
- Formato da história: "Como [ator], quero [ação] para [benefício]"
- Cada US deve ter entre 3 e 6 critérios de aceitação objetivos
- Identifique dependências entre US (qual US precisa estar pronta antes)
- Agrupe as US por épico
- Ordene as US respeitando a sequência técnica de dependências
- NÃO inclua estimativas de dias (serão coletadas do desenvolvedor)

Retorne APENAS um JSON válido, sem texto adicional, sem markdown, sem backticks:

{{
  "user_stories": [
    {{
      "id": "US-001",
      "epico_id": "EP01",
      "rf_ids": ["RF01"],
      "titulo": "título curto da US",
      "historia": "Como [ator], quero [ação] para [benefício].",
      "tela_contexto": "nome da tela ou contexto no sistema",
      "campos_elementos": ["campo ou elemento 1", "campo ou elemento 2"],
      "criterios_aceitacao": ["critério 1", "critério 2"],
      "dependencias": ["US-00X — descrição da dependência ou 'Nenhuma'"],
      "prioridade": "Alta"
    }}
  ]
}}

REQUISITOS EXTRAÍDOS:
{requisitos}
"""

PROMPT_MUDANCAS_BASE = """
Você é um analista de requisitos sênior especializado em Scrum e projetos legados.

Abaixo está uma lista de mudanças, melhorias e ajustes identificados em um sistema já existente,
sem documento formal de requisitos. Sua tarefa é organizar essas mudanças em um backlog estruturado.

Regras:
- Agrupe mudanças relacionadas em Épicos temáticos (ex: "Melhoria de UX", "Correção de Fluxo Cadastral")
- Trate cada mudança descrita como um ou mais Requisitos Funcionais (RF)
- Identifique atores afetados (ex: Usuário, Administrador, Sistema)
- Se não houver informação suficiente para um campo, use valores genéricos como "Sistema legado" ou "A definir"
- NÃO invente mudanças além das descritas — só organize o que foi informado

Retorne APENAS um JSON válido, sem texto adicional, sem markdown, sem backticks:

{{
  "projeto": "nome do projeto ou sistema",
  "sistema": "nome do sistema/plataforma",
  "descricao": "resumo das mudanças em 2-3 linhas",
  "atores": ["lista de atores afetados"],
  "epicos": [
    {{
      "id": "EP01",
      "nome": "nome do épico",
      "descricao": "descrição do épico"
    }}
  ],
  "requisitos_funcionais": [
    {{
      "id": "RF01",
      "descricao": "descrição completa da mudança/melhoria",
      "epico_id": "EP01",
      "atores": ["Usuário"]
    }}
  ],
  "regras_negocio": [],
  "requisitos_nao_funcionais": [],
  "fora_de_escopo": [],
  "dependencias_tecnicas": [],
  "observacoes": "observações relevantes sobre as mudanças"
}}

MUDANÇAS IDENTIFICADAS:
{mudancas}
"""

def chamar_claude(prompt: str) -> str:
    """Chama a API da Anthropic e retorna o texto da resposta."""
    response = CLIENT.messages.create(
        model=MODEL,
        max_tokens=16000,
        messages=[{"role": "user", "content": prompt}]
    )
    return response.content[0].text
def extrair_json(texto: str) -> dict:
    """Extrai e parseia JSON da resposta do modelo com fallback robusto."""
    if not texto or not texto.strip():
        raise ValueError(
            "O modelo retornou uma resposta vazia.\n"
            "Possíveis causas: documento muito longo, modelo sobrecarregado.\n"
            "Tente novamente — o modelo às vezes falha na primeira tentativa."
        )

    texto = texto.strip()
    # Remove blocos de código markdown se existirem
    texto = re.sub(r"```json\s*", "", texto)
    texto = re.sub(r"```\s*", "", texto)
    texto = texto.strip()

    # Encontra o primeiro { e o último }
    inicio = texto.find("{")
    fim    = texto.rfind("}") + 1

    if inicio == -1 or fim == 0:
        preview = texto[:400] if len(texto) > 400 else texto
        raise ValueError(
            f"Nenhum JSON encontrado na resposta.\n"
            f"O modelo respondeu:\n{preview}\n\n"
            f"Tente rodar novamente — modelos locais ocasionalmente fogem do formato."
        )

    try:
        return json.loads(texto[inicio:fim])
    except json.JSONDecodeError as e:
        # Tenta corrigir JSON truncado adicionando fechamento
        trecho = texto[inicio:fim]
        preview = trecho[:400] if len(trecho) > 400 else trecho
        raise ValueError(
            f"JSON inválido na resposta do modelo: {e}\n"
            f"Trecho recebido:\n{preview}\n\n"
            f"O modelo pode ter truncado a resposta. Tente rodar novamente."
        )

# ════════════════════════════════════════════════════════════
# COLETA INTERATIVA DE DIAS ÚTEIS
# ════════════════════════════════════════════════════════════

def ask_int(texto):
    while True:
        try:
            val = int(pergunta(texto))
            if val >= 0: return val
            aviso("Digite um valor zero ou positivo.")
        except ValueError:
            aviso("Digite apenas um número inteiro.")

def coletar_horas_uteis(user_stories: list, horas_por_sprint: int) -> list:
    """Pergunta ao DEV as horas úteis de dev/teste para cada US e retorna a lista atualizada."""
    secao("Estimativa de Horas Úteis por User Story")
    info("Para cada US, informe quantas horas úteis levará para Dev e Teste/Homologação.")
    info(f"Regra: recomendável não ultrapassar {horas_por_sprint} horas úteis (capacidade efetiva da sprint).")
    info(f"Se uma US exigir mais de {horas_por_sprint} horas, será sugerida a divisão automaticamente.\n")

    for us in user_stories:
        print(f"\n  {BOLD}{us['id']}{RESET} — {us['titulo']}")
        print(f"  {CINZA}{us['historia']}{RESET}")
        if us.get("dependencias") and us["dependencias"] != ["Nenhuma"]:
            print(f"  {CINZA}Depende de: {', '.join(us['dependencias'])}{RESET}")

        while True:
            h_dev    = ask_int(f"  Horas úteis de Desenvolvimento para {us['id']}:")
            h_teste  = ask_int(f"  Horas úteis de Testes/Homologação para {us['id']}:")
            
            total = h_dev + h_teste
            
            if 0 < total <= horas_por_sprint:
                us["horas_dev"] = h_dev
                us["horas_teste"] = h_teste
                us["horas_estimadas"] = total
                ok(f"{us['id']} → Dev: {h_dev}h | Teste: {h_teste}h | Total: {total} hora(s)")
                break
            elif total > horas_por_sprint:
                aviso(f"Total de {total} horas úteis excede a capacidade efetiva do sprint ({horas_por_sprint} horas).")
                confirmar = pergunta("Manter assim mesmo e transbordar o sprint? (s/n):").lower()
                if confirmar == "s":
                    us["horas_dev"] = h_dev
                    us["horas_teste"] = h_teste
                    us["horas_estimadas"] = total
                    ok(f"{us['id']} → Total: {total} horas (sprint estendido)")
                    break
            else:
                aviso("O total de horas não pode ser zero.")

    return user_stories

# ════════════════════════════════════════════════════════════
# DISTRIBUIÇÃO EM SPRINTS SEMANAIS
        # ════════════════════════════════════════════════════════════

def distribuir_sprints(user_stories: list, horas_por_sprint: int, horas_por_semana: int, horas_deploy_sprint: int) -> list:
    """
    Distribui as US em sprints baseando-se nas horas definidas.
    Respeita dependências e agrupa US possíveis combinando suas estimativas.
    Retorna lista de sprints com US alocadas.
    """
    # Mapa de US por ID para lookup de dependências
    us_map = {us["id"]: us for us in user_stories}
    us_concluidas = set()
    sprints = []
    sprint_num = 1
    semana = 1
    us_pendentes = list(user_stories)  # cópia

    while us_pendentes:
        # US disponíveis = sem dependências pendentes
        disponiveis = []
        for us in us_pendentes:
            deps_ids = []
            for dep in us.get("dependencias", []):
                match = re.search(r"US-\d+", dep)
                if match:
                    deps_ids.append(match.group())
            if all(d in us_concluidas for d in deps_ids):
                disponiveis.append(us)

        if not disponiveis:
            # Dependências circulares ou bloqueio — força a primeira US pendente
            aviso(f"Dependência não resolvida detectada. Alocando {us_pendentes[0]['id']} mesmo assim.")
            disponiveis = [us_pendentes[0]]

        # Tenta preencher o sprint com até horas_por_sprint úteis
        sprint_cards = []
        horas_usadas = 0
        LIMITE = horas_por_sprint

        for us in disponiveis:
            h = us.get("horas_estimadas", 3)
            # Agrupa se couber no sprint limite
            if horas_usadas == 0:
                sprint_cards.append(us)
                horas_usadas += h
                us_pendentes.remove(us)
                if horas_usadas >= LIMITE:
                    break
            elif horas_usadas + h <= LIMITE:
                sprint_cards.append(us)
                horas_usadas += h
                us_pendentes.remove(us)
                if horas_usadas >= LIMITE:
                    break

        # Cálculo de semanas estendidas se a tarefa for maior que o sprint
        semanas_estimadas = max(1, (horas_usadas + horas_por_semana - 1) // horas_por_semana)
        
        if semanas_estimadas == 1:
            semana_str = f"Sem. {semana}"
        else:
            semana_str = f"Sem. {semana}–{semana + semanas_estimadas - 1}"

        # Entrega resumida
        entrega = " + ".join(us["titulo"] for us in sprint_cards)

        # Cálculo dos totais parciais
        tot_dev = sum(us.get("horas_dev", 0) for us in sprint_cards)
        tot_teste = sum(us.get("horas_teste", 0) for us in sprint_cards)

        # RFs cobertos
        rfs = []
        for us in sprint_cards:
            rfs.extend(us.get("rf_ids", []))
        rfs_str = ", ".join(sorted(set(rfs)))

        sprints.append({
            "numero": sprint_num,
            "semana": semana_str,
            "cards": sprint_cards,
            "tot_dev": tot_dev,
            "tot_teste": tot_teste,
            "horas_deploy": horas_deploy_sprint, # Fixo por sprint
            "tot_horas_us": horas_usadas,
            "horas_estimadas": horas_usadas + horas_deploy_sprint, # Horas Totais do Sprint
            "entrega": entrega,
            "rfs": rfs_str,
        })

        for us in sprint_cards:
            us_concluidas.add(us["id"])

        sprint_num += 1
        semana += semanas_estimadas

    return sprints

# ════════════════════════════════════════════════════════════
# GERAÇÃO DO DOCX
# ════════════════════════════════════════════════════════════

def add_heading(doc, text, level=1, color="1F4E79"):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    run = p.add_run(text)
    run.bold = True
    run.font.size = Pt(18 if level == 1 else 14)
    run.font.color.rgb = RGBColor.from_string(color)
    run.font.name = "Arial"
    p.paragraph_format.space_before = Pt(16)
    p.paragraph_format.space_after = Pt(6)
    return p

def add_par(doc, text, size=10, bold=False, color="222222", italic=False):
    p = doc.add_paragraph()
    run = p.add_run(text)
    run.font.size = Pt(size)
    run.font.bold = bold
    run.font.italic = italic
    run.font.color.rgb = RGBColor.from_string(color)
    run.font.name = "Arial"
    p.paragraph_format.space_before = Pt(3)
    p.paragraph_format.space_after = Pt(3)
    return p

def add_bullet(doc, text, size=10):
    p = doc.add_paragraph(style="List Bullet")
    run = p.add_run(text)
    run.font.size = Pt(size)
    run.font.name = "Arial"
    return p

def set_cell(cell, text, bg="FFFFFF", bold=False, size=9, color="222222", center=False):
    cell.text = ""
    p = cell.paragraphs[0]
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER if center else WD_ALIGN_PARAGRAPH.LEFT
    run = p.add_run(str(text))
    run.font.size = Pt(size)
    run.font.bold = bold
    run.font.name = "Arial"
    run.font.color.rgb = RGBColor.from_string(color)
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:val"), "clear")
    shd.set(qn("w:color"), "auto")
    shd.set(qn("w:fill"), bg)
    tcPr.append(shd)

def gerar_docx(projeto: dict, requisitos: dict, user_stories: list, sprints: list, caminho: str):
    doc = DocxDocument()
    # Margens
    for section in doc.sections:
        section.left_margin   = Inches(0.9)
        section.right_margin  = Inches(0.9)
        section.top_margin    = Inches(0.8)
        section.bottom_margin = Inches(0.8)

    # Capa
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run(requisitos.get("sistema", projeto).upper())
    run.bold = True; run.font.size = Pt(26); run.font.name = "Arial"
    run.font.color.rgb = RGBColor.from_string("1F4E79")

    p2 = doc.add_paragraph()
    p2.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run2 = p2.add_run(f"BACKLOG — FASE 1  |  Versão 1.0  |  {datetime.date.today().strftime('%d/%m/%Y')}")
    run2.font.size = Pt(12); run2.font.name = "Arial"
    run2.font.color.rgb = RGBColor.from_string("888888")

    add_par(doc, f"{len(sprints)} sprints semanais  |  {len(user_stories)} User Stories  |  Scrum", size=11, color="2E75B6", bold=True)
    doc.add_page_break()

    # Descrição do projeto
    add_heading(doc, "1. Visão Geral do Projeto")
    add_par(doc, requisitos.get("descricao", ""))
    add_par(doc, f"Atores: {', '.join(requisitos.get('atores', []))}", bold=True, color="1F4E79")

    if requisitos.get("fora_de_escopo"):
        add_par(doc, "Fora de escopo nesta fase:", bold=True, color="C55A11")
        for item in requisitos["fora_de_escopo"]:
            add_bullet(doc, item)

    doc.add_page_break()

    # Épicos
    add_heading(doc, "2. Épicos")
    t = doc.add_table(rows=1, cols=3)
    t.style = "Table Grid"
    hdr = t.rows[0].cells
    for i, h in enumerate(["ID", "Épico", "Descrição"]):
        set_cell(hdr[i], h, bg="1F4E79", bold=True, color="FFFFFF", center=True)
    for i, ep in enumerate(requisitos.get("epicos", [])):
        row = t.add_row().cells
        bg = "EBF3FB" if i % 2 == 0 else "F5F9FF"
        set_cell(row[0], ep["id"],       bg=bg, bold=True, center=True)
        set_cell(row[1], ep["nome"],     bg=bg)
        set_cell(row[2], ep["descricao"],bg=bg)
    doc.add_paragraph()

    # Sprints
    add_heading(doc, "3. Sprints e User Stories")
    doc.add_page_break()

    for sprint in sprints:
        add_heading(doc, f"Sprint {sprint['numero']} — {sprint['semana']}", level=2, color="2E75B6")

        info_lines = [
            f"Objetivo: {sprint['entrega']}",
            f"Semana: {sprint['semana']}  |  Dev: {sprint.get('tot_dev', 0)}h  |  Teste: {sprint.get('tot_teste', 0)}h  |  Deploy (Sprint): {sprint.get('horas_deploy', 0)}h  |  Total: {sprint['horas_estimadas']}h",
            f"RFs cobertos: {sprint['rfs']}",
        ]
        for line in info_lines:
            add_par(doc, line, size=9, color="444444")
        doc.add_paragraph()

        for us in sprint["cards"]:
            # Header da US
            t2 = doc.add_table(rows=1, cols=5)
            t2.style = "Table Grid"
            cells = t2.rows[0].cells
            set_cell(cells[0], us["id"],                    bg="1F4E79", bold=True, color="FFFFFF", center=True, size=10)
            set_cell(cells[1], us["titulo"],                bg="2E75B6", bold=True, color="FFFFFF", size=9)
            set_cell(cells[2], f"Épico: {us['epico_id']}", bg="1F4E79", bold=True, color="FFFFFF", center=True, size=8)
            set_cell(cells[3], f"RF: {', '.join(us.get('rf_ids',[]))}",  bg="375623", bold=True, color="FFFFFF", center=True, size=8)
            horas_str = f"Dev:{us.get('horas_dev','?')}h | Tst:{us.get('horas_teste','?')}h | Tot:{us.get('horas_estimadas','?')}h"
            set_cell(cells[4], horas_str, bg="C55A11", bold=True, color="FFFFFF", center=True, size=7)

            # Corpo da US — duas colunas
            t3 = doc.add_table(rows=1, cols=2)
            t3.style = "Table Grid"
            cl, cr = t3.rows[0].cells

            # Coluna esquerda
            cl.text = ""
            for label, content in [
                ("História:", us["historia"]),
                ("Critérios de Aceitação:", ""),
            ]:
                p = cl.add_paragraph()
                r = p.add_run(label); r.bold = True; r.font.size = Pt(9); r.font.name = "Arial"
                r.font.color.rgb = RGBColor.from_string("1F4E79")
                if content:
                    p2 = cl.add_paragraph()
                    r2 = p2.add_run(content); r2.font.size = Pt(9)
                    r2.font.name = "Arial"; r2.font.italic = True

            for crit in us.get("criterios_aceitacao", []):
                pb = cl.add_paragraph(style="List Bullet")
                rb = pb.add_run(crit); rb.font.size = Pt(8); rb.font.name = "Arial"

            p_dep = cl.add_paragraph()
            rd = p_dep.add_run(f"Dependência: "); rd.bold = True; rd.font.size = Pt(8); rd.font.name = "Arial"
            rd.font.color.rgb = RGBColor.from_string("666666")
            rd2 = p_dep.add_run(", ".join(us.get("dependencias", ["Nenhuma"])))
            rd2.font.size = Pt(8); rd2.font.name = "Arial"; rd2.font.italic = True

            p_prio = cl.add_paragraph()
            rp = p_prio.add_run(f"Prioridade: {us.get('prioridade','Alta')}")
            rp.bold = True; rp.font.size = Pt(8); rp.font.name = "Arial"
            cor_prio = "C00000" if us.get("prioridade") == "Alta" else "C55A11"
            rp.font.color.rgb = RGBColor.from_string(cor_prio)

            # Coluna direita
            cr.text = ""
            p_tela = cr.add_paragraph()
            rt = p_tela.add_run("Tela / Contexto:"); rt.bold = True
            rt.font.size = Pt(9); rt.font.name = "Arial"
            rt.font.color.rgb = RGBColor.from_string("1F4E79")
            p_tela2 = cr.add_paragraph()
            rt2 = p_tela2.add_run(us.get("tela_contexto", ""))
            rt2.font.size = Pt(9); rt2.font.name = "Arial"

            p_campos = cr.add_paragraph()
            rc = p_campos.add_run("Campos / Elementos:"); rc.bold = True
            rc.font.size = Pt(9); rc.font.name = "Arial"
            rc.font.color.rgb = RGBColor.from_string("1F4E79")

            for campo in us.get("campos_elementos", []):
                pb2 = cr.add_paragraph(style="List Bullet")
                rb2 = pb2.add_run(campo); rb2.font.size = Pt(8); rb2.font.name = "Arial"

            # Shading das células
            for cell, bg in [(cl, "FAFCFF"), (cr, "F0F6FF")]:
                tc = cell._tc; tcPr = tc.get_or_add_tcPr()
                shd = OxmlElement("w:shd")
                shd.set(qn("w:val"), "clear"); shd.set(qn("w:color"), "auto")
                shd.set(qn("w:fill"), bg); tcPr.append(shd)

            doc.add_paragraph()

    doc.save(caminho)

# ════════════════════════════════════════════════════════════
# GERAÇÃO DA PLANILHA
# ════════════════════════════════════════════════════════════

def fill_xl(h): return PatternFill("solid", start_color=h, fgColor=h)
def fnt_xl(bold=False,size=10,color="222222"):
    return Font(bold=bold,size=size,color=color,name="Arial")
def aln_xl(h="left",v="center",wrap=False):
    return Alignment(horizontal=h,vertical=v,wrap_text=wrap)
def brd_xl():
    s=Side(style="thin",color="BBBBBB")
    return Border(left=s,right=s,top=s,bottom=s)

MODULO_CORES = [
    "DEEAF1","EBF3FB","E8F5E3","F3E8FF","FFF4EC","E8EAF6",
    "DDEEFF","FFE8D6","E8FFE8","F0E8FF","FFEEDD","E0F0FF",
]

def gerar_xlsx(requisitos: dict, user_stories: list, sprints: list, configs: dict, caminho: str):
    wb = openpyxl.Workbook()

    projeto_nome = requisitos.get("projeto", "Projeto")

    # ── Aba Resumo Executivo ─────────────────────────────────
    ws0 = wb.active; ws0.title = "Resumo"
    ws0.merge_cells("A1:D1")
    c0 = ws0["A1"]; c0.value = f"RESUMO EXECUTIVO · {projeto_nome.upper()}"
    c0.font = fnt_xl(bold=True, size=12, color="FFFFFF"); c0.fill = fill_xl("1F4E79")
    c0.alignment = aln_xl(h="center", v="center"); ws0.row_dimensions[1].height = 28
    
    total_us = len(user_stories)
    total_horas = sum(us.get("horas_estimadas", 0) for us in user_stories) + (len(sprints) * configs["horas_deploy_sprint"])
    dias_uteis_eq = total_horas / configs["horas_por_dia"] if configs.get("horas_por_dia") else 0

    items = [
        ("Configuração da Sprint", ""),
        ("Duração da Sprint:", f"{configs['semanas_sprint']} Semana(s)"),
        ("Horas Produtivas por Dev/Dia:", f"{configs['horas_por_dia']} horas"),
        ("Orçamento de Deploy por Sprint:", f"{configs['horas_deploy_sprint']} horas"),
        ("Capacidade Bruta da Sprint:", f"{configs['horas_por_sprint']} horas"),
        ("", ""),
        ("Estimativas do Projeto", ""),
        ("Total de Sprints Estimados:", f"{len(sprints)} Sprints"),
        ("Total de User Stories:", f"{total_us} US"),
        ("Esforço Global Estimado:", f"{total_horas} horas"),
        ("Equivalente em Dias de Trabalho:", f"~{dias_uteis_eq:,.1f} Dias Úteis"),
    ]
    
    ws0.column_dimensions["A"].width = 35
    ws0.column_dimensions["B"].width = 25
    
    for r, (label, val) in enumerate(items, 3):
        c_label = ws0.cell(row=r, column=1, value=label)
        c_val = ws0.cell(row=r, column=2, value=val)
        if label in ["Configuração da Sprint", "Estimativas do Projeto"]:
            c_label.font = fnt_xl(bold=True, color="FFFFFF")
            c_label.fill = fill_xl("2E75B6"); c_val.fill = fill_xl("2E75B6")
        else:
            c_label.font = fnt_xl(bold=True); c_label.fill = fill_xl("EBF3FB")
            c_val.font = fnt_xl(); c_val.fill = fill_xl("F5F9FF")
        if label:
            c_label.border = brd_xl(); c_val.border = brd_xl()
        ws0.row_dimensions[r].height = 22

    # ── Aba Backlog ──────────────────────────────────────────
    ws = wb.create_sheet("Backlog")
    ws.freeze_panes = "A3"

    ws.merge_cells("A1:M1")
    c = ws["A1"]
    c.value = f"BACKLOG · {projeto_nome.upper()}  |  {len(sprints)} sprints semanais  |  {len(user_stories)} User Stories"
    c.font = fnt_xl(bold=True,size=11,color="FFFFFF"); c.fill = fill_xl("1F4E79")
    c.alignment = aln_xl(h="center",v="center"); ws.row_dimensions[1].height = 28

    COLS = ["Sprint","Semana","ID","Título","Épico","RF","Horas Dev","Horas Teste","Total Horas","Prioridade","Status","Observações","Responsável"]
    WIDS = [9,10,9,44,9,14,10,10,10,12,20,30,18]
    for i,(col,w) in enumerate(zip(COLS,WIDS),1):
        cell = ws.cell(row=2,column=i,value=col)
        cell.font=fnt_xl(bold=True,color="FFFFFF"); cell.fill=fill_xl("2E75B6")
        cell.alignment=aln_xl(h="center",v="center",wrap=True); cell.border=brd_xl()
        ws.column_dimensions[get_column_letter(i)].width=w
    ws.row_dimensions[2].height=20

    row_i = 3
    for sp_idx, sprint in enumerate(sprints):
        bg = MODULO_CORES[sp_idx % len(MODULO_CORES)]
        for us in sprint["cards"]:
            vals = [
                f"S{sprint['numero']:02d}", sprint["semana"], us["id"], us["titulo"],
                us.get("epico_id",""), ", ".join(us.get("rf_ids",[])),
                us.get("horas_dev",""), us.get("horas_teste",""),
                us.get("horas_estimadas",""), us.get("prioridade","Alta"),
            ]
            for col_i, val in enumerate(vals, 1):
                cell = ws.cell(row=row_i, column=col_i, value=val)
                cell.fill=fill_xl(bg); cell.border=brd_xl()
                # Ajusta alinhamento para o Título (col 4)
                cell.alignment=aln_xl(h="center" if col_i != 4 else "left",v="center",wrap=True)
                cell.font=fnt_xl(bold=(col_i in [1,3]),size=9,
                    color="C00000" if (col_i==11 and val=="Alta") else
                          "C55A11" if (col_i==11 and val=="Média") else "222222")
            # Status dropdown (col 11 agora, Obs 12, Resp 13)
            sc = ws.cell(row=row_i, column=11, value="To Do")
            sc.font=fnt_xl(size=9,color="555555"); sc.fill=fill_xl("F5F5F5")
            sc.alignment=aln_xl(h="center",v="center"); sc.border=brd_xl()
            ws.cell(row=row_i,column=12,value="").border=brd_xl()
            ws.cell(row=row_i,column=13,value="").border=brd_xl()
            ws.row_dimensions[row_i].height=26; row_i+=1

    dv=DataValidation(type="list",formula1='"To Do,In Progress,In Review,Done,Blocked"',showDropDown=True)
    dv.sqref=f"K3:K{row_i-1}"; ws.add_data_validation(dv)

    # ── Aba Sprints ──────────────────────────────────────────
    ws2 = wb.create_sheet("Sprints"); ws2.freeze_panes="A3"
    ws2.merge_cells("A1:H1")
    c2=ws2["A1"]; c2.value=f"SPRINTS · {projeto_nome.upper()}"
    c2.font=fnt_xl(bold=True,size=11,color="FFFFFF"); c2.fill=fill_xl("1F4E79")
    c2.alignment=aln_xl(h="center",v="center"); ws2.row_dimensions[1].height=28

    COLS2=["Sprint","Semana","Card(s)","Entrega Principal","Horas Dev","Horas Teste","Horas Deploy","Total Horas","RFs","Status"]
    WIDS2=[10,12,20,44,10,10,10,10,22,20]
    for i,(c,w) in enumerate(zip(COLS2,WIDS2),1):
        cell=ws2.cell(row=2,column=i,value=c)
        cell.font=fnt_xl(bold=True,color="FFFFFF"); cell.fill=fill_xl("2E75B6")
        cell.alignment=aln_xl(h="center",v="center",wrap=True); cell.border=brd_xl()
        ws2.column_dimensions[get_column_letter(i)].width=w
    ws2.row_dimensions[2].height=20

    for sp_idx, sprint in enumerate(sprints, 3):
        bg=MODULO_CORES[(sp_idx-3) % len(MODULO_CORES)]
        cards_str=" + ".join(us["id"] for us in sprint["cards"])
        vals=[f"Sprint {sprint['numero']:02d}", sprint["semana"], cards_str,
              sprint["entrega"], sprint.get("tot_dev",0), sprint.get("tot_teste",0),
              sprint.get("horas_deploy",0), sprint["horas_estimadas"], sprint["rfs"]]
        for col_i,val in enumerate(vals,1):
            cell=ws2.cell(row=sp_idx,column=col_i,value=val)
            cell.fill=fill_xl(bg); cell.border=brd_xl()
            cell.alignment=aln_xl(h="center" if col_i!=4 else "left",v="center",wrap=True)
            cell.font=fnt_xl(size=9,bold=(col_i==1))
        
        # Formula automatizada de Status da Sprint pegando o status na aba do Backlog
        # Se TODAS daquela sprint forem Done -> Done. Se TODAS forem To Do -> To Do. Senao In Progress.
        str_sp = f"S{sprint['numero']:02d}"
        f_status = f'=IF(COUNTIFS(Backlog!A:A,"{str_sp}",Backlog!K:K,"Done")=COUNTIFS(Backlog!A:A,"{str_sp}"),"Done",IF(COUNTIFS(Backlog!A:A,"{str_sp}",Backlog!K:K,"To Do")=COUNTIFS(Backlog!A:A,"{str_sp}"),"To Do","In Progress"))'
        sc=ws2.cell(row=sp_idx,column=10,value=f_status)
        sc.font=fnt_xl(size=9,color="555555"); sc.fill=fill_xl("F5F5F5")
        sc.alignment=aln_xl(h="center",v="center"); sc.border=brd_xl()
        ws2.row_dimensions[sp_idx].height=26

    # Total
    tr=3+len(sprints)
    ws2.merge_cells(f"A{tr}:D{tr}") # Merge A-D
    tc=ws2[f"A{tr}"]; tc.value=f"TOTAL — {len(sprints)} Sprints"
    tc.font=fnt_xl(bold=True,size=10,color="FFFFFF"); tc.fill=fill_xl("1F4E79")
    tc.alignment=aln_xl(h="center",v="center"); tc.border=brd_xl()
    
    # Preencher células do Total E F G H I
    for col_i,val in enumerate([f"=SUM(E3:E{tr-1})", f"=SUM(F3:F{tr-1})", f"=SUM(G3:G{tr-1})", f"=SUM(H3:H{tr-1})", "Total Final",""],5):
        cell=ws2.cell(row=tr,column=col_i,value=val)
        cell.font=fnt_xl(bold=True,color="FFFFFF"); cell.fill=fill_xl("1F4E79")
        cell.alignment=aln_xl(h="center",v="center"); cell.border=brd_xl()
    ws2.row_dimensions[tr].height=26

    # ── Aba Kanban ───────────────────────────────────────────
    ws3=wb.create_sheet("Kanban"); ws3.merge_cells("A1:E1")
    ck=ws3["A1"]; ck.value="KANBAN · Mova os cards conforme o andamento"
    ck.font=fnt_xl(bold=True,size=11,color="FFFFFF"); ck.fill=fill_xl("1F4E79")
    ck.alignment=aln_xl(h="center",v="center"); ws3.row_dimensions[1].height=28
    KCOLS=["To Do","In Progress","In Review","Done","Blocked"]
    KCOLORS=["F5F5F5","FFF9C4","FFF3E0","E8F5E9","FFEBEE"]
    for i,(col,w) in enumerate(zip(KCOLS,[30,30,30,30,30]),1):
        cell=ws3.cell(row=2,column=i,value=col)
        cell.font=fnt_xl(bold=True,color="FFFFFF"); cell.fill=fill_xl("2E75B6")
        cell.alignment=aln_xl(h="center",v="center"); cell.border=brd_xl()
        ws3.column_dimensions[get_column_letter(i)].width=w
    ws3.row_dimensions[2].height=22
    # Preencher com fórmula =IF linha a linha refletindo o Backlog
    for row_i, us in enumerate(user_stories, 3):
        for col_i, stat in enumerate(KCOLS, 1):
            f = f'=IF(Backlog!$K{row_i}="{stat}", Backlog!$C{row_i} & " · " & Backlog!$D{row_i}, "")'
            c = ws3.cell(row=row_i, column=col_i, value=f)
            c.font = fnt_xl(size=9)
            c.alignment = aln_xl(v="center", wrap=True)
            c.fill = fill_xl(KCOLORS[col_i-1])
            c.border = brd_xl()
        ws3.row_dimensions[row_i].height = 40
    # ── Aba Legenda ──────────────────────────────────────────
    ws4=wb.create_sheet("Legenda")
    for col,w in zip(["A","B","C","D"],[22,40,22,30]):
        ws4.column_dimensions[col].width=w
    ws4.merge_cells("A1:D1")
    cl=ws4["A1"]; cl.value="LEGENDA"
    cl.font=fnt_xl(bold=True,size=12,color="FFFFFF"); cl.fill=fill_xl("1F4E79")
    cl.alignment=aln_xl(h="center",v="center"); ws4.row_dimensions[1].height=28
    hdr_items=[("A3","STATUS","B3","Significado"),("C3","PRIORIDADE","D3","Significado")]
    for a,ta,b,tb in hdr_items:
        for ref,val in [(a,ta),(b,tb)]:
            c=ws4[ref]; c.value=val
            c.font=fnt_xl(bold=True,color="FFFFFF"); c.fill=fill_xl("2E75B6")
            c.alignment=aln_xl(h="center",v="center"); c.border=brd_xl()
    status_items=[("To Do","Card ainda não iniciado","F5F5F5"),
                  ("In Progress","Em desenvolvimento neste sprint","FFF9C4"),
                  ("In Review","Aguardando validação / aprovação","FFF3E0"),
                  ("Done","Entregue, testado e validado","E8F5E9"),
                  ("Blocked","Impedido por dependência ou dúvida","FFEBEE")]
    prio_items=[("Alta","Bloqueia o sprint se não entregue","FFEBEE","C00000"),
                ("Média","Importante mas não bloqueia","FFF3E0","C55A11")]
    for i,(s,d,cor) in enumerate(status_items,4):
        for col,val,align in [(1,s,"center"),(2,d,"left")]:
            c=ws4.cell(row=i,column=col,value=val)
            c.font=fnt_xl(bold=(col==1),size=10); c.fill=fill_xl(cor)
            c.alignment=aln_xl(h=align,v="center"); c.border=brd_xl()
        ws4.row_dimensions[i].height=22
    for i,(p,d,cor,fc) in enumerate(prio_items,4):
        for col,val,align,fc2 in [(3,p,"center",fc),(4,d,"left","222222")]:
            c=ws4.cell(row=i,column=col,value=val)
            c.font=fnt_xl(bold=(col==3),size=10,color=fc2); c.fill=fill_xl(cor)
            c.alignment=aln_xl(h=align,v="center"); c.border=brd_xl()

    wb.save(caminho)

# ════════════════════════════════════════════════════════════
# FLUXO PRINCIPAL
# ════════════════════════════════════════════════════════════

def main():
    titulo("AGENTE DE BACKLOG SCRUM")
    info("Este agente gera o backlog com sprints semanais, User Stories e planilha de acompanhamento.")
    info(f"Modelo: {MODEL} via API Anthropic\n")

    # ── Duração da Sprint ────────────────────────────────────────
    secao("Configuração da Sprint")
    info("O padrão de mercado é entre 1 e 4 semanas.")
    semanas_sprint = ask_int("Duração da sprint em semanas (1 a 4):")
    if semanas_sprint < 1: semanas_sprint = 1
    
    horas_por_dia = ask_int("Horas produtivas por dia do desenvolvedor (ex: 6 ou 8):")
    if horas_por_dia < 1: horas_por_dia = 8
    
    # Cada sprint assume N horas produtivas por semana
    horas_por_semana = horas_por_dia * 5
    horas_por_sprint = semanas_sprint * horas_por_semana
    
    horas_deploy_sprint = ask_int("Horas reservadas para Deploy/Produção por sprint (ex: 4 ou 8):")
    if horas_deploy_sprint < 0: horas_deploy_sprint = 0
    horas_disp_us = max(0, horas_por_sprint - horas_deploy_sprint)
    
    ok(f"Sprints de {horas_por_sprint}h: ({horas_disp_us}h para Histórias + {horas_deploy_sprint}h de Deploy).\n")

    # ── Seleção de modo ──────────────────────────────────────────
    secao("Selecione o Modo de Operação")
    info("  [1] Tenho documento(s) de requisitos  → backlog completo (DOCX + XLSX)")
    info("  [2] Vou descrever mudanças/melhorias  → planilha de acompanhamento (XLSX)")
    while True:
        modo = pergunta("Modo (1 ou 2):").strip()
        if modo in ("1", "2"):
            break
        aviso("Opção inválida. Digite 1 ou 2.")

    # ── Coleta de entrada ────────────────────────────────────────
    if modo == "1":
        texto_entrada = coletar_documentos()
        prompt_extracao = PROMPT_EXTRACAO_BASE.format(documentos=texto_entrada)
    else:
        titulo("MODO: MUDANÇAS / MELHORIAS")
        info("Use este modo para projetos legados sem documento de requisitos formal.")
        info("Descreva as mudanças — o agente organiza em épicos, US e sprints.\n")
        texto_entrada = coletar_mudancas()
        prompt_extracao = PROMPT_MUDANCAS_BASE.format(mudancas=texto_entrada)

    # ── Extração / organização via IA ────────────────────────────
    secao("Analisando com IA...")
    print("   Aguarde — isso pode levar alguns segundos...")
    resposta_ext = chamar_claude(prompt_extracao)
    requisitos = extrair_json(resposta_ext)
    ok(f"Projeto identificado: {requisitos.get('projeto','—')}")
    ok(f"Épicos extraídos: {len(requisitos.get('epicos',[]))}")
    ok(f"Requisitos funcionais: {len(requisitos.get('requisitos_funcionais',[]))}")
    if requisitos.get("regras_negocio"):
        ok(f"Regras de negócio: {len(requisitos.get('regras_negocio',[]))}")
    ok(f"Atores: {', '.join(requisitos.get('atores',[]))}")

    # Mostrar épicos para confirmação
    secao("Épicos identificados:")
    for ep in requisitos.get("epicos", []):
        info(f"  {ep['id']} — {ep['nome']}")

    confirmar = pergunta("\nOs épicos estão corretos? (s/n):").lower()
    if confirmar != "s":
        if modo == "1":
            aviso("Ajuste o documento de requisitos e rode o agente novamente.")
            aviso("Dica: adicione mais contexto ou separe os épicos manualmente no documento.")
        else:
            aviso("Ajuste as descrições das mudanças e rode o agente novamente.")
        return

    # ── Gerar User Stories ───────────────────────────────────────
    secao("Gerando User Stories...")
    prompt_us = PROMPT_GERAR_US_BASE.format(requisitos=json.dumps(requisitos, ensure_ascii=False))
    resposta_us = chamar_claude(prompt_us)
    dados_us = extrair_json(resposta_us)
    user_stories = dados_us.get("user_stories", [])
    ok(f"User Stories geradas: {len(user_stories)}")

    secao("User Stories geradas:")
    for us in user_stories:
        info(f"  {us['id']} — {us['titulo']}  [{us.get('prioridade','Alta')}]")

    confirmar2 = pergunta("\nAs User Stories parecem corretas? (s/n):").lower()
    if confirmar2 != "s":
        aviso("Ajuste as informações e rode novamente, ou edite manualmente a planilha gerada.")

    # ── Coletar estimativas (Horas) ──────────────────────────────
    user_stories = coletar_horas_uteis(user_stories, horas_disp_us)

    # ── Distribuir sprints ───────────────────────────────────────
    secao("Distribuindo em Sprints...")
    sprints = distribuir_sprints(user_stories, horas_disp_us, horas_por_semana, horas_deploy_sprint)
    ok(f"Total de sprints gerados: {len(sprints)}")

    secao("Distribuição final:")
    for sp in sprints:
        cards_str = " + ".join(us["id"] for us in sp["cards"])
        info(f"  Sprint {sp['numero']:02d} ({sp['semana']}) → {cards_str} [Total: {sp['horas_estimadas']}h]")

    # ── Configuração de saída ────────────────────────────────────
    secao("Configuração de Saída")
    nome_projeto = pergunta("Nome do projeto para nomear os arquivos (sem espaços, ex: SGAU_FASE2):") or "projeto"
    nome_projeto = re.sub(r"[^a-zA-Z0-9_\-]", "_", nome_projeto)

    pasta_saida = pergunta("Pasta de saída (ENTER para ./backlog):").strip() or "./backlog"
    Path(pasta_saida).mkdir(parents=True, exist_ok=True)

    data_str = datetime.date.today().strftime("%Y%m%d")
    caminho_xlsx = f"{pasta_saida}/Acompanhamento_{nome_projeto}_{data_str}.xlsx"

    # ── Gerar arquivos ───────────────────────────────────────────
    secao("Gerando arquivos...")

    if modo == "1":
        caminho_docx = f"{pasta_saida}/Backlog_{nome_projeto}_{data_str}.docx"
        gerar_docx(nome_projeto, requisitos, user_stories, sprints, caminho_docx)
        ok(f"Backlog DOCX: {caminho_docx}")

    configs = {
        "semanas_sprint": semanas_sprint,
        "horas_por_dia": horas_por_dia,
        "horas_por_semana": horas_por_semana,
        "horas_por_sprint": horas_por_sprint,
        "horas_deploy_sprint": horas_deploy_sprint,
    }

    gerar_xlsx(requisitos, user_stories, sprints, configs, caminho_xlsx)
    ok(f"Planilha XLSX: {caminho_xlsx}")

    # ── Resumo final ─────────────────────────────────────────────
    titulo("BACKLOG GERADO COM SUCESSO!")
    if modo == "1":
        print(f"""
  {VERDE}Arquivos gerados:{RESET}
    📄  {caminho_docx}
    📊  {caminho_xlsx}

  {CINZA}Próximos passos:{RESET}
    1. Valide o backlog com o analista de requisitos
    2. Apresente as US ao solicitante na reunião de alinhamento
    3. Ajuste dias úteis se necessário e redistribua manualmente
""")
    else:
        print(f"""
  {VERDE}Arquivo gerado:{RESET}
    📊  {caminho_xlsx}

  {CINZA}Próximos passos:{RESET}
    1. Revise os cards na aba Backlog e ajuste prioridades
    2. Atualize o status dos cards conforme o andamento (aba Kanban)
    3. Compartilhe com a equipe e alinhe a sequência de sprints
""")

if __name__ == "__main__":
    main()
